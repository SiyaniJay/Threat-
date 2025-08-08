# main_copy.py

import os
from email import policy
from email.parser import BytesParser
from bs4 import BeautifulSoup
import spacy
from docx import Document
from datetime import datetime

from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.text_rank import TextRankSummarizer

# Import the function for the security chatbot
from chatbot import get_featherless_response as get_sec_bot_response

# — Load spaCy model —
nlp = spacy.load("en_core_web_md")

def parse_eml(uploaded_file):
    return BytesParser(policy=policy.default).parse(uploaded_file)

def extract_email_parts(msg):
    subject = msg["subject"]
    sender = msg["from"]
    body = ""
    if msg.is_multipart():
        for part in msg.walk():
            ctype = part.get_content_type()
            if ctype == "text/plain":
                body = part.get_content()
                break
            elif ctype == "text/html":
                html = part.get_content()
                body = BeautifulSoup(html, "html.parser").get_text()
                break
    else:
        body = msg.get_content()
    return {"subject": subject, "from": sender, "body_text": body.strip()}

def get_named_entities(text):
    doc = nlp(text)
    return [ent.text for ent in doc.ents]

def get_relevance_score(text):
    doc = nlp(text)
    reference = nlp("Critical, Cobalt Strike, cybersecurity, threat, affecting university systems, students, and campus infrastructure.")
    return doc.similarity(reference)

def summarize_email(text, sentence_count=3):
    parser = PlaintextParser.from_string(text, Tokenizer("english"))
    summarizer = TextRankSummarizer()
    summary = summarizer(parser.document, sentence_count)
    return " ".join(str(sentence) for sentence in summary)

def classify_urgency(text, keywords, similarity_score):
    text_lower = text.lower()
    if any(k in text_lower for k in keywords):
        return "Red"
    if similarity_score > 0.9:
        return "Red"
    if similarity_score > 0.6:
        return "Orange"
    return "Yellow"

def generate_report(data, output_dir="output"):
    os.makedirs(output_dir, exist_ok=True)
    doc = Document()
    doc.add_heading("Threat Report", 0)
    doc.add_paragraph(datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    doc.add_paragraph(f"Subject: {data['subject']}")
    doc.add_paragraph(f"From: {data['from']}")
    doc.add_paragraph(f"Urgency: {data['urgency']}")
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(data["summary"])
    doc.add_heading("Named Entities", level=1)
    doc.add_paragraph(", ".join(data["entities"]))
    doc.add_heading("Full Email Body", level=1)
    doc.add_paragraph(data["body_text"])

    filename = f"{output_dir}/ThreatReport_{data['urgency'].upper()}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(filename)
    return filename

# --- Bridge function for the security chatbot ---
def ask_sec_chatbot(question):
    """
    A bridge function to pass a question to the FoundationSec chatbot.
    """
    return get_sec_bot_response(question)
